import streamlit as st
from docx import Document
from io import BytesIO

def export_report(df, student):
    doc=Document()
    doc.add_heading(f"Отчет: {student}", level=1)
    doc.add_paragraph("...здесь содержимое отчета...")
    buf=BytesIO()
    doc.save(buf)
    st.download_button("Скачать отчет DOCX", buf.getvalue(), file_name=f"{student}_report.docx")
